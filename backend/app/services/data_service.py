import os
import re
import pandas as pd
from datetime import datetime, date
from typing import Optional, Tuple, Dict, Any, List
from app.config import settings

TEACHING_PLAN_COLUMNS = ["session_number", "planned_date", "module", "topic_title", "subtopics"]
CANDIDATE_COLUMNS = ["candidate_id", "candidate_name", "email", "phone", "batch_id", "status"]
BATCH_COLUMNS = ["batch_id", "batch_name", "program_name", "start_date", "end_date", "status", "coordinator_name"]
ADMIN_COLUMNS = ["admin_id", "admin_name", "admin_email", "admin_phone", "role"]

MODULE_HEADING_RE = re.compile(r"Module\s+\d+\s*:\s*(.+)", re.IGNORECASE)

def align_columns(df: pd.DataFrame, expected_columns: List[str], dataset_name: str) -> Tuple[pd.DataFrame, bool]:
    df.columns = [str(c).strip().lower().replace(" ", "_") for c in df.columns]
    missing = [c for c in expected_columns if c not in df.columns]
    extra = [c for c in df.columns if c not in expected_columns]
    if missing:
        print(f"WARNING [{dataset_name}]: missing expected column(s) {missing}. Found: {list(df.columns)}")
    if extra:
        df = df.drop(columns=extra)
    return df, (len(missing) == 0)

def read_docx_table(path: str) -> pd.DataFrame:
    from docx import Document  # type: ignore[import-untyped]
    doc = Document(path)
    if not doc.tables:
        raise ValueError("No table found in the .docx file. Please include the data as a Word table.")
    table = doc.tables[0]
    rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
    header, *data = rows
    return pd.DataFrame(data, columns=header)

def read_pdf_table(path: str) -> pd.DataFrame:
    import pdfplumber
    all_rows = []
    header = None
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            for t in page.extract_tables():
                if not t:
                    continue
                if header is None:
                    header, *rows = t
                else:
                    rows = t[1:] if t[0] == header else t
                all_rows.extend(rows)
    if not all_rows:
        raise ValueError("No table detected in the PDF.")
    return pd.DataFrame(all_rows, columns=header)

def read_txt_table(path: str) -> pd.DataFrame:
    for sep in [",", "\t", ";", "|"]:
        try:
            df = pd.read_csv(path, sep=sep)
            if df.shape[1] > 1:
                return df
        except Exception:
            continue
    return pd.read_csv(path, sep=None, engine="python")

def read_tabular_file(path: str) -> pd.DataFrame:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".csv":
        return pd.read_csv(path)
    elif ext in (".xlsx", ".xls"):
        return pd.read_excel(path)
    elif ext == ".txt":
        return read_txt_table(path)
    elif ext == ".docx":
        return read_docx_table(path)
    elif ext == ".pdf":
        return read_pdf_table(path)
    else:
        raise ValueError(f"Unsupported file type: {ext}. Please upload .csv, .xlsx, .txt, .docx, or .pdf")

def read_teaching_plan_pdf(path: str) -> pd.DataFrame:
    """
    Parses a teaching-plan PDF where each module's sessions sit under a bold section heading
    (e.g. 'Module 1: Foundations of Agentic AI') followed by a ruled table with columns
    [Session No., Date, Topic Title, Sub-Topics].
    """
    import pdfplumber
    rows = []
    current_module = None

    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            words = page.extract_words()
            lines_by_top = {}
            for w in words:
                top = round(w["top"])
                lines_by_top.setdefault(top, []).append(w)

            heading_positions = []
            for top in sorted(lines_by_top.keys()):
                ws = sorted(lines_by_top[top], key=lambda w: w["x0"])
                line_text = " ".join(w["text"] for w in ws)
                m = MODULE_HEADING_RE.search(line_text)
                if m:
                    heading_positions.append((top, m.group(1).strip()))

            for table in page.find_tables():
                table_top = table.bbox[1]
                headings_above = [h for h in heading_positions if h[0] < table_top]
                if headings_above:
                    current_module = headings_above[-1][1]

                data = table.extract()
                if not data:
                    continue
                header, *data_rows = data

                def clean(cell):
                    return " ".join(str(cell).split()) if cell is not None else ""

                for row in data_rows:
                    if not row or row[0] is None:
                        continue
                    session_no_raw = str(row[0]).strip()
                    if not session_no_raw.isdigit():
                        continue

                    rows.append({
                        "session_number": int(session_no_raw),
                        "planned_date": clean(row[1]) if len(row) > 1 else "",
                        "module": current_module or "",
                        "topic_title": clean(row[2]) if len(row) > 2 else "",
                        "subtopics": clean(row[3]) if len(row) > 3 else "",
                    })

    df = pd.DataFrame(rows, columns=["session_number", "planned_date", "module", "topic_title", "subtopics"])
    if not df.empty:
        df = df.sort_values("session_number").reset_index(drop=True)
    return df

class DataService:
    def __init__(self):
        self.data_dir = settings.DATA_DIR

    def get_candidates(self) -> pd.DataFrame:
        csv_path = os.path.join(self.data_dir, "candidates.csv")
        xlsx_path = os.path.join(self.data_dir, "candidates.xlsx")
        
        if os.path.exists(csv_path):
            df = pd.read_csv(csv_path)
        elif os.path.exists(xlsx_path):
            df = pd.read_excel(xlsx_path)
            df, _ = align_columns(df, CANDIDATE_COLUMNS, "Candidates")
            df.to_csv(csv_path, index=False)
        else:
            return pd.DataFrame(columns=CANDIDATE_COLUMNS)
        
        df, _ = align_columns(df, CANDIDATE_COLUMNS, "Candidates")
        # Ensure candidate_id is numeric if possible
        try:
            df["candidate_id"] = pd.to_numeric(df["candidate_id"])
        except Exception:
            pass
        return df

    def save_candidates(self, df: pd.DataFrame) -> bool:
        df, ok = align_columns(df, CANDIDATE_COLUMNS, "Candidates")
        if ok:
            csv_path = os.path.join(self.data_dir, "candidates.csv")
            df.to_csv(csv_path, index=False)
            return True
        return False

    def add_candidate(self, row_dict: Dict[str, Any]) -> bool:
        df = self.get_candidates()
        new_id = 1 if df.empty else df["candidate_id"].max() + 1
        row_dict["candidate_id"] = new_id
        new_row = pd.DataFrame([row_dict])
        df = pd.concat([df, new_row], ignore_index=True)
        return self.save_candidates(df)

    def delete_candidates(self, candidate_ids: List[Any]) -> bool:
        df = self.get_candidates()
        if not candidate_ids:
            return True
            
        candidate_ids_str = [str(c) for c in candidate_ids]
        
        try:
            from app.tools.attendance_tools import load_attendance_log, ATTENDANCE_CSV
            att_log = load_attendance_log()
            if not att_log.empty:
                att_log = att_log[~att_log["candidate_id"].astype(str).isin(candidate_ids_str)]
                att_log.to_csv(ATTENDANCE_CSV, index=False)
        except Exception:
            pass
            
        try:
            from app.tools.communication_tools import load_comm_log, COMM_LOG_CSV
            comm_log = load_comm_log()
            if not comm_log.empty:
                comm_log = comm_log[~comm_log["candidate_id"].astype(str).isin(candidate_ids_str)]
                comm_log.to_csv(COMM_LOG_CSV, index=False)
        except Exception:
            pass

        df = df[~df["candidate_id"].astype(str).isin(candidate_ids_str)]
        ok = self.save_candidates(df)
        self.clean_orphan_logs()
        return ok

    def parse_candidates_file(self, file_path: str) -> Tuple[pd.DataFrame, bool, str]:
        try:
            df = read_tabular_file(file_path)
            df, ok = align_columns(df, CANDIDATE_COLUMNS, "Candidates")
            if not ok or df.empty:
                return df, False, f"Parsed file is missing required columns: {CANDIDATE_COLUMNS}"
            return df, True, f"Successfully parsed {len(df)} candidates."
        except Exception as e:
            return pd.DataFrame(), False, f"Failed to parse candidates file: {str(e)}"

    def get_batches(self) -> pd.DataFrame:
        csv_path = os.path.join(self.data_dir, "batches.csv")
        xlsx_path = os.path.join(self.data_dir, "batches.xlsx")
        
        if os.path.exists(csv_path):
            df = pd.read_csv(csv_path)
        elif os.path.exists(xlsx_path):
            df = pd.read_excel(xlsx_path)
            df, _ = align_columns(df, BATCH_COLUMNS, "Batches")
            df.to_csv(csv_path, index=False)
        else:
            return pd.DataFrame(columns=BATCH_COLUMNS)
            
        df, _ = align_columns(df, BATCH_COLUMNS, "Batches")
        return df

    def save_batches(self, df: pd.DataFrame) -> bool:
        df, ok = align_columns(df, BATCH_COLUMNS, "Batches")
        if ok:
            csv_path = os.path.join(self.data_dir, "batches.csv")
            df.to_csv(csv_path, index=False)
            return True
        return False

    def add_batch(self, row_dict: Dict[str, Any]) -> bool:
        df = self.get_batches()
        new_row = pd.DataFrame([row_dict])
        df = pd.concat([df, new_row], ignore_index=True)
        return self.save_batches(df)

    def delete_batches(self, batch_ids: List[Any]) -> bool:
        df = self.get_batches()
        if not batch_ids:
            return True
            
        batch_ids_str = [str(b) for b in batch_ids]
        
        c_df = self.get_candidates()
        if not c_df.empty:
            cands_to_delete = c_df[c_df["batch_id"].astype(str).isin(batch_ids_str)]["candidate_id"].tolist()
            if cands_to_delete:
                self.delete_candidates(cands_to_delete)
                
        df = df[~df["batch_id"].astype(str).isin(batch_ids_str)]
        ok = self.save_batches(df)
        self.clean_orphan_logs()
        return ok

    def parse_batches_file(self, file_path: str) -> Tuple[pd.DataFrame, bool, str]:
        try:
            df = read_tabular_file(file_path)
            df, ok = align_columns(df, BATCH_COLUMNS, "Batches")
            if not ok or df.empty:
                return df, False, f"Parsed file is missing required columns: {BATCH_COLUMNS}"
            return df, True, f"Successfully parsed {len(df)} batches."
        except Exception as e:
            return pd.DataFrame(), False, f"Failed to parse batches file: {str(e)}"
    def clean_orphan_logs(self):
        try:
            from app.tools.attendance_tools import load_attendance_log, ATTENDANCE_CSV
            att_log = load_attendance_log()
            if not att_log.empty:
                c_df = self.get_candidates()
                valid_cand_ids = set(c_df["candidate_id"].astype(str).tolist()) if not c_df.empty else set()
                tp = self.get_teaching_plan()
                valid_sessions = set(tp["session_number"].astype(int).tolist()) if not tp.empty else set()
                
                filtered_att = att_log[
                    att_log["candidate_id"].astype(str).isin(valid_cand_ids) &
                    att_log["session_number"].astype(int).isin(valid_sessions)
                ]
                filtered_att.to_csv(ATTENDANCE_CSV, index=False)
        except Exception:
            pass

        try:
            from app.tools.communication_tools import load_comm_log, COMM_LOG_CSV
            comm_log = load_comm_log()
            if not comm_log.empty:
                c_df = self.get_candidates()
                valid_cand_ids = set(c_df["candidate_id"].astype(str).tolist()) if not c_df.empty else set()
                tp = self.get_teaching_plan()
                valid_sessions = set(tp["session_number"].astype(int).tolist()) if not tp.empty else set()
                
                filtered_comm = comm_log[
                    comm_log["candidate_id"].astype(str).isin(valid_cand_ids) &
                    comm_log["session_number"].astype(int).isin(valid_sessions)
                ]
                filtered_comm.to_csv(COMM_LOG_CSV, index=False)
        except Exception:
            pass

    def delete_sessions(self, session_numbers: List[Any]) -> bool:
        df = self.get_teaching_plan()
        if not session_numbers:
            return True
        session_nums_int = []
        for s in session_numbers:
            try:
                session_nums_int.append(int(s))
            except Exception:
                pass
        df = df[~df["session_number"].astype(int).isin(session_nums_int)]
        ok = self.save_teaching_plan(df)
        self.clean_orphan_logs()
        return ok

    def get_admin_details(self) -> pd.DataFrame:
        csv_path = os.path.join(self.data_dir, "admin_details.csv")
        xlsx_path = os.path.join(self.data_dir, "admin_details.xlsx")
        
        if os.path.exists(csv_path):
            df = pd.read_csv(csv_path)
        elif os.path.exists(xlsx_path):
            df = pd.read_excel(xlsx_path)
            df, _ = align_columns(df, ADMIN_COLUMNS, "Admin Details")
            df.to_csv(csv_path, index=False)
        else:
            return pd.DataFrame(columns=ADMIN_COLUMNS)
            
        df, _ = align_columns(df, ADMIN_COLUMNS, "Admin Details")
        return df

    def get_teaching_plan(self) -> pd.DataFrame:
        csv_path = os.path.join(self.data_dir, "teaching_plan.csv")
        if os.path.exists(csv_path):
            df = pd.read_csv(csv_path)
            df, _ = align_columns(df, TEACHING_PLAN_COLUMNS, "Teaching Plan")
            return df
        return pd.DataFrame(columns=TEACHING_PLAN_COLUMNS)

    def save_teaching_plan(self, df: pd.DataFrame) -> bool:
        df, ok = align_columns(df, TEACHING_PLAN_COLUMNS, "Teaching Plan")
        if ok:
            csv_path = os.path.join(self.data_dir, "teaching_plan.csv")
            df.to_csv(csv_path, index=False)
            return True
        return False

    def add_session(self, row_dict: Dict[str, Any]) -> bool:
        df = self.get_teaching_plan()
        new_row = pd.DataFrame([row_dict])
        df = pd.concat([df, new_row], ignore_index=True)
        return self.save_teaching_plan(df)

    def parse_and_save_teaching_plan_file(self, file_path: str) -> Tuple[pd.DataFrame, bool, str]:
        ext = os.path.splitext(file_path)[1].lower()
        df = pd.DataFrame()
        try:
            if ext == ".pdf":
                df = read_teaching_plan_pdf(file_path)
                if df.empty:
                    df = read_pdf_table(file_path)
            else:
                df = read_tabular_file(file_path)
            
            df, ok = align_columns(df, TEACHING_PLAN_COLUMNS, "Teaching Plan")
            if not ok or df.empty:
                return df, False, f"Parsed file is missing required columns: {TEACHING_PLAN_COLUMNS}"
                
            self.save_teaching_plan(df)
            return df, True, f"Successfully parsed and saved {len(df)} sessions to Teaching Plan."
        except Exception as e:
            return df, False, f"Failed to parse teaching plan: {str(e)}"

    def get_today_date(self) -> date:
        if settings.SIMULATED_TODAY:
            try:
                return datetime.strptime(settings.SIMULATED_TODAY, "%d-%m-%Y").date()
            except ValueError:
                pass
        return date.today()

    def get_today_session(self) -> Optional[Dict[str, Any]]:
        tp = self.get_teaching_plan()
        if tp.empty:
            return None
        today_str = self.get_today_date().strftime("%d-%m-%Y")
        match = tp[tp["planned_date"].astype(str) == today_str]
        if match.empty:
            # If no session matching exact date, return the next upcoming or closest session
            # for demo/coordinator convenience
            return None
        return match.iloc[0].to_dict()

    def get_session_details(self, session_number: int) -> Dict[str, Any]:
        tp = self.get_teaching_plan()
        default_session = {
            "session_number": session_number,
            "planned_date": "N/A",
            "module": "Unknown",
            "topic_title": "Unknown",
            "subtopics": ""
        }
        if tp.empty:
            return default_session
        row = tp[tp["session_number"] == int(session_number)]
        if row.empty:
            return default_session
        return row.iloc[0].to_dict()

data_service = DataService()
